import sys
import random
from PySide6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                           QHBoxLayout, QLabel, QSpinBox, QPushButton, 
                           QTableWidget, QTableWidgetItem, QMessageBox,
                           QGroupBox, QFileDialog, QScrollArea, QCheckBox,
                           QTabWidget, QComboBox, QHeaderView, QFormLayout,
                           QLineEdit, QInputDialog, QButtonGroup, QRadioButton)
from PySide6.QtCore import Qt
import os
from datetime import datetime
import tempfile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn, OxmlElement
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from docx2pdf import convert
from docx import Document
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PySide6.QtGui import QFont
from docx.shared import Inches, Pt;
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import openpyxl
from models import Role,PatientProfile,Preference
from openpyxl.worksheet.datavalidation import DataValidation
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from models import Base
from models import MealItem,ExclusionRule,HealthCondition
from sqlalchemy.exc import IntegrityError
from login import LoginDialog
def create_element(name):
    return OxmlElement(name)

def create_attribute(element,name,value):
    element.set(qn(name),value)
def create_dropdown_element(options, selected=None):
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    ddl = OxmlElement('w:dropDownList')
    for option in options:
        li = OxmlElement('w:listItem')
        li.set(qn('w:displayText'), option)
        li.set(qn('w:value'), option)
        ddl.append(li)
    sdtPr.append(ddl)
    sdt.append(sdtPr)

    sdtContent = OxmlElement('w:sdtContent')
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = selected or (options[0] if options else '')
    r.append(t)
    p.append(r)
    sdtContent.append(p)
    sdt.append(sdtContent)
    return sdt

class MealPlanner(QMainWindow):
    RATING_WEIGHTS = {
    "أحبه كثيراً":    3.0,
    "أحبه بشكل متوسط": 2.0,
    "أحبه قليلاً":    1.0,
    "لا أحبه":        0.5,
    "Not Rated":      3.0,
    }
    def __init__(self,db_session,current_user):
        super().__init__()
        self.db = db_session
        self.user = current_user
        self.setWindowTitle("Meal Planner")
        self.setMinimumSize(1200, 800)
        self.health_conditions = []
        
        # Set application font for Arabic text
        app_font = QFont("Arial")
        app_font.setPointSize(11)
        QApplication.setFont(app_font)
        
        # Set application-wide style
        self.setStyleSheet("""
            QMainWindow {
                background-color: #ffffff;
            }
            QGroupBox {
                background-color: #ffffff;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
                margin-top: 1.5em;
                padding: 15px;
                font-weight: bold;
                color: #111827;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 10px;
                color: #111827;
                font-size: 14px;
                background-color: #ffffff;
            }
            QPushButton {
                background-color: #3b82f6;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 8px;
                font-weight: bold;
                font-size: 13px;
                min-width: 120px;
            }
            QPushButton:hover {
                background-color: #2563eb;
            }
            QPushButton:disabled {
                background-color: #9ca3af;
                color: #f3f4f6;
            }
            QSpinBox {
                padding: 8px;
                border: 2px solid #e5e7eb;
                border-radius: 8px;
                background-color: #ffffff;
                color: #111827;
                font-size: 13px;
                min-width: 80px;
            }
            QSpinBox:hover {
                border-color: #3b82f6;
            }
            QSpinBox::up-button, QSpinBox::down-button {
                border: none;
                background-color: #f3f4f6;
                border-radius: 4px;
                margin: 1px;
            }
            QSpinBox::up-button:hover, QSpinBox::down-button:hover {
                background-color: #e5e7eb;
            }
            QTableWidget {
                background-color: #ffffff;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
                gridline-color: #f3f4f6;
                color: #111827;
                font-size: 13px;
            }
            QTableWidget::item {
                padding: 12px;
                border-bottom: 1px solid #f3f4f6;
            }
            QTableWidget::item:selected {
                background-color: #dbeafe;
                color: #1e40af;
            }
            QHeaderView::section {
                background-color: #f9fafb;
                padding: 12px;
                border: none;
                border-bottom: 2px solid #e5e7eb;
                color: #111827;
                font-weight: bold;
                font-size: 13px;
            }
            QTabWidget::pane {
                border: 1px solid #e5e7eb;
                border-radius: 12px;
                background-color: #ffffff;
                top: -1px;
            }
            QTabBar::tab {
                background-color: #f9fafb;
                border: 1px solid #e5e7eb;
                padding: 10px 20px;
                margin-right: 4px;
                border-top-left-radius: 8px;
                border-top-right-radius: 8px;
                color: #6b7280;
                font-size: 13px;
            }
            QTabBar::tab:selected {
                background-color: #ffffff;
                border-bottom: none;
                color: #3b82f6;
                font-weight: bold;
            }
            QTabBar::tab:hover:!selected {
                color: #3b82f6;
            }
            QScrollArea {
                border: none;
                background-color: #ffffff;
            }
            QCheckBox {
                spacing: 8px;
                color: #111827;
                font-size: 13px;
            }
            QCheckBox::indicator {
                width: 20px;
                height: 20px;
                border-radius: 6px;
            }
            QCheckBox::indicator:unchecked {
                border: 2px solid #e5e7eb;
                background-color: #ffffff;
            }
            QCheckBox::indicator:checked {
                background-color: #3b82f6;
                border: 2px solid #3b82f6;
            }
            QCheckBox::indicator:hover {
                border-color: #3b82f6;
            }
            QLabel {
                color: #111827;
                font-size: 13px;
                font-weight: 500;
            }
            QComboBox {
                padding: 8px 12px;
                border: 2px solid #e5e7eb;
                border-radius: 8px;
                background-color: #ffffff;
                color: #111827;
                min-width: 250px;
                font-size: 13px;
                selection-background-color: #dbeafe;
                selection-color: #1e40af;
                text-align: right;
            }
            QComboBox:hover {
                border-color: #3b82f6;
            }
            QComboBox::drop-down {
                border: none;
                width: 24px;
            }
            QComboBox::down-arrow {
                width: 12px;
                height: 12px;
                margin-right: 8px;
                image: none;
                border: none;
                background: none;
            }
            QComboBox::down-arrow:after {
                content: "";
                display: block;
                width: 0;
                height: 0;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 5px solid #6b7280;
            }
            QComboBox QAbstractItemView {
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                background-color: #ffffff;
                selection-background-color: #dbeafe;
                selection-color: #1e40af;
                padding: 4px;
            }
            QScrollArea, QWidget#scrollContent {
                background-color: #ffffff;
                border: none;
            }
            QScrollArea {
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
        """)
        
        # Initialize the main widget and layout
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        layout = QVBoxLayout(main_widget)
        layout.setSpacing(20)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Create tab widget with modern styling
        tabs = QTabWidget()
        tabs.setStyleSheet("""
            QTabWidget::pane {
                border: 1px solid #e0e0e0;
                border-radius: 8px;
                background-color: white;
            }
            QTabBar::tab {
                background-color: #f5f5f5;
                border: 1px solid #e0e0e0;
                padding: 8px 16px;
                margin-right: 2px;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
            }
            QTabBar::tab:selected {
                background-color: white;
                border-bottom: none;
            }
        """)
        ratings = [
            ("أحبه كثيراً", "☺️", "😊", "#22c55e"),
            ("أحبه بشكل متوسط", "🙂", "😀", "#0ea5e9"),
            ("أحبه قليلاً", "😐", "🙁", "#f59e0b"),
            ("لا أحبه", "😕", "😞", "#ef4444")
        ]
        self.ratings = ratings
        self.RATING_LABELS = [text for text, outlined, filled, color in ratings]
        self.rating_map = { label: i for i, label in enumerate(self.RATING_LABELS) }
        # Create main tab
        main_tab = QWidget()
        main_layout = QVBoxLayout(main_tab)
        #health tabls
        health_tab = QWidget()
        health_layout = QVBoxLayout(health_tab)
        health_group = QGroupBox("Health Conditions")
        health_group_layout = QVBoxLayout()
        self.healthy_checkbox = QCheckBox("Healthy (1)")
        self.diabetes_checkbox = QCheckBox("Diabetes (2)")
        self.kidney_checkbox = QCheckBox("Kidney Disease (3)")
        # Connect health condition checkboxes to update method
        self.healthy_checkbox.stateChanged.connect(self.update_health_conditions)
        self.diabetes_checkbox.stateChanged.connect(self.update_health_conditions)
        self.kidney_checkbox.stateChanged.connect(self.update_health_conditions)
        # add checkboxes to health group
        health_group_layout.addWidget(self.healthy_checkbox)
        health_group_layout.addWidget(self.diabetes_checkbox)
        health_group_layout.addWidget(self.kidney_checkbox)
        health_group.setLayout(health_group_layout)
        # Add health group to health tab
        health_layout.addWidget(health_group)
        health_layout.addStretch()
        # Create input section
        input_group = QGroupBox("Category Settings")
        input_layout = QHBoxLayout()
        
        # Category A input
        category_a_layout = QVBoxLayout()
        category_a_label = QLabel("Category A Count:")
        self.category_a_spin = QSpinBox()
        self.category_a_spin.setRange(1, 6)
        self.category_a_spin.setValue(4)
        category_a_layout.addWidget(category_a_label)
        category_a_layout.addWidget(self.category_a_spin)
        
        # Category B input
        category_b_layout = QVBoxLayout()
        category_b_label = QLabel("Category B Count:")
        self.category_b_spin = QSpinBox()
        self.category_b_spin.setRange(1, 6)
        self.category_b_spin.setValue(2)
        category_b_layout.addWidget(category_b_label)
        category_b_layout.addWidget(self.category_b_spin)
        
        # Category C input
        category_c_layout = QVBoxLayout()
        category_c_label = QLabel("Category C Count:")
        self.category_c_spin = QSpinBox()
        self.category_c_spin.setRange(1, 6)
        self.category_c_spin.setValue(1)
        category_c_layout.addWidget(category_c_label)
        category_c_layout.addWidget(self.category_c_spin)
        
        input_layout.addLayout(category_a_layout)
        input_layout.addLayout(category_b_layout)
        input_layout.addLayout(category_c_layout)
        input_group.setLayout(input_layout)
        
        # Update button layout with spacing
        button_layout = QHBoxLayout()
        button_layout.setSpacing(10)
        self.generate_button = QPushButton("Generate Meal Plan")
        self.generate_button.setStyleSheet("""
            QPushButton {
                background-color: #10b981;
                color: white;
                border: none;
                padding: 12px 24px;
                border-radius: 8px;
                font-weight: bold;
                font-size: 14px;
                min-width: 150px;
            }
            QPushButton:hover {
                background-color: #059669;
            }
        """)
        self.generate_button.clicked.connect(self.generate_meal_plan)
        self.save_excel_button = QPushButton("Save to Excel")
        self.save_excel_button.clicked.connect(self.save_to_excel)
        self.save_excel_button.setEnabled(False)
        self.export_to_word_template_button = QPushButton("Export to Word Template")
        self.export_to_word_template_button.clicked.connect(self.save_to_template_word)
        self.export_to_word_template_button.setEnabled(False)
        self.save_word_button = QPushButton("Save to Word")
        self.save_word_button.clicked.connect(self.save_to_word)
        self.save_word_button.setEnabled(False)
        self.save_pdf_button = QPushButton("Save to PDF")
        self.save_pdf_button.clicked.connect(self.save_to_pdf)
        self.save_pdf_button.setEnabled(False)
        self.save_gdocs_button = QPushButton("Save to Google Docs")
        self.save_gdocs_button.clicked.connect(self.save_to_gdoc)
        self.save_gdocs_button.setEnabled(False)
        button_layout.addWidget(self.generate_button)
        button_layout.addWidget(self.save_excel_button)
        button_layout.addWidget(self.save_word_button)
        button_layout.addWidget(self.save_pdf_button)
        button_layout.addWidget(self.save_gdocs_button)
        button_layout.addWidget(self.export_to_word_template_button)
        
        # Table widget for displaying meal plan
        table_group = QGroupBox("Meal Plan")
        table_layout = QVBoxLayout()
        self.table = QTableWidget()
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(["اليوم", "الإفطار", "الغداء", "العشاء"])
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.setStyleSheet("""
            QTableWidget {
                background-color: white;
                border: 1px solid #e0e0e0;
                border-radius: 8px;
                gridline-color: #e0e0e0;
            }
            QTableWidget::item {
                padding: 8px;
            }
            QHeaderView::section {
                background-color: #f5f5f5;
                padding: 8px;
                border: none;
                border-bottom: 1px solid #e0e0e0;
            }
        """)
        table_layout.addWidget(self.table)
        table_group.setLayout(table_layout)
        self.items = self.db.query(MealItem).all()
        # Add widgets to main tab layout
        main_layout.addWidget(input_group)
        main_layout.addLayout(button_layout)
        main_layout.addWidget(table_group)
        
        # Create exclusion tab
        exclusion_tab = QWidget()
        exclusion_layout = QVBoxLayout(exclusion_tab)
        
        # Create scroll area for exclusions
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll_content = QWidget()
        scroll_content.setObjectName("scrollContent")
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setSpacing(10)
        scroll_layout.setContentsMargins(15, 15, 15, 15)
        
        # Create checkboxes for each meal item
        self.exclusion_checkboxes = {}
        for m in self.items:
            checkbox = QCheckBox(m.name)
        # for item in MEAL_ITEMS:
        #     checkbox = QCheckBox(item["name"])
            checkbox.setLayoutDirection(Qt.RightToLeft)
            checkbox.setChecked(False)
            # self.exclusion_checkboxes[item["name"]] = checkbox
            self.exclusion_checkboxes[m.name] = checkbox
            scroll_layout.addWidget(checkbox)
        
        scroll.setWidget(scroll_content)
        exclusion_layout.addWidget(scroll)
        
        # Add tabs to tab widget
        tabs.addTab(main_tab, "Meal Planner")
        tabs.addTab(health_tab, "Health Conditions")
        tabs.addTab(exclusion_tab, "Exclude Items")
        
        # Create and add preferences tab
        preferences_tab = QWidget()
        preferences_layout = QVBoxLayout(preferences_tab)
        
        # Add title/info section with modern styling
        info_group = QGroupBox()
        info_group.setStyleSheet("""
            QGroupBox {
                background-color: #ffffff;
                border: 1px solid #e5e7eb;
                border-radius: 16px;
                padding: 20px;
                margin-bottom: 20px;
                box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
            }
        """)
        info_layout = QHBoxLayout(info_group)
        
        name_label = QLabel("الاسم:")
        name_label.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        name_input = QLineEdit()
        self.name_input = name_input
        save_pref_button = QPushButton("Save Preferences")
        save_pref_button.setStyleSheet("""
            QPushButton {
                background-color: #3b82f6;
                color: white;
                border: none;
                padding: 12px 24px;
                border-radius: 8px;
                font-weight: bold;
                font-size: 14px;
                min-width: 150px;
            }
            QPushButton:hover {
                background-color: #2563eb;
            }
        """)
        save_pref_button.clicked.connect(self.save_preferences)
        name_input.setAlignment(Qt.AlignRight)
        name_input.setPlaceholderText("ادخل الاسم")
        name_input.setStyleSheet("""
            QLineEdit {
                padding: 12px;
                border: 2px solid #e5e7eb;
                border-radius: 12px;
                font-size: 14px;
                background-color: #f8fafc;
            }
            QLineEdit:focus {
                border-color: #3b82f6;
                background-color: white;
            }
        """)
        
        sample_label = QLabel("رقم العينة:")
        sample_label.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        sample_input = QLineEdit()
        self.sample_input = sample_input
        self.setup_sample_input_listener()
        sample_input.setAlignment(Qt.AlignRight)
        sample_input.setPlaceholderText("ادخل رقم العينة")
        sample_input.setStyleSheet("""
            QLineEdit {
                padding: 12px;
                border: 2px solid #e5e7eb;
                border-radius: 12px;
                font-size: 14px;
                background-color: #f8fafc;
            }
            QLineEdit:focus {
                border-color: #3b82f6;
                background-color: white;
            }
        """)
        
        info_layout.addWidget(sample_input)
        info_layout.addWidget(sample_label)
        info_layout.addWidget(name_input)
        info_layout.addWidget(name_label)
        
        preferences_layout.addWidget(info_group)
        preferences_layout.addWidget(save_pref_button)
        # Create scroll area for the form
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll_content = QWidget()
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setSpacing(20)
        
        # Define rating options with emojis and colors
        ratings = [
            ("أحبه كثيراً", "☺️", "😊", "#22c55e"),    # Green
            ("أحبه بشكل متوسط", "🙂", "😀", "#0ea5e9"), # Blue
            ("أحبه قليلاً", "😐", "🙁", "#f59e0b"),      # Yellow
            ("لا أحبه", "😕", "😞", "#ef4444")          # Red
        ]
        
        # Add headers section with modern design
        headers_widget = QWidget()
        headers_layout = QHBoxLayout(headers_widget)
        headers_layout.setContentsMargins(0, 0, 0, 0)
        
        ratings_widget = QWidget()
        ratings_layout = QHBoxLayout(ratings_widget)
        ratings_layout.setContentsMargins(0, 0, 0, 0)
        ratings_layout.setSpacing(30)  # Increased spacing between ratings
        
        for rating_text, outlined, filled, color in ratings:
            label = QLabel(f"{rating_text}\n{outlined}")
            label.setStyleSheet(f"""
                font-weight: bold;
                color: {color};
                font-size: 14px;
                padding: 8px;
                border-radius: 8px;
                background-color: #f8fafc;
            """)
            label.setAlignment(Qt.AlignCenter)
            ratings_layout.addWidget(label)
        
        headers_layout.addWidget(ratings_widget)
        
        scroll_layout.addWidget(headers_widget)
        
        # Group meals by eat_time with modern card design
        meal_times = ["Breakfast", "Lunch", "Dinner"]
        meal_headers = {
            "Breakfast": "خيارات الفطور",
            "Lunch": "خيارات الغداء",
            "Dinner": "خيارات العشاء"
        }
        self.preference_buttons = {}
        for meal_time in meal_times:
            # Add section header with modern styling
            section_group = QGroupBox()
            section_group.setStyleSheet("""
                QGroupBox {
                    background-color: white;
                    border: 1px solid #e5e7eb;
                    border-radius: 16px;
                    padding: 20px;
                    margin-top: 10px;
                    box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
                }
            """)
            section_layout = QVBoxLayout(section_group)
            
            section_label = QLabel(meal_headers[meal_time])
            section_label.setStyleSheet("""
                font-weight: bold;
                color: #1e293b;
                font-size: 16px;
                padding: 12px;
                border-bottom: 2px solid #e5e7eb;
                margin-bottom: 10px;
            """)
            section_label.setAlignment(Qt.AlignRight)
            section_layout.addWidget(section_label)
            
            # Get meals for this section
            meals = [meal for meal in self.items if meal.eat_time == meal_time]
            
            for meal in meals:
                row_widget = QWidget()
                row_widget.setStyleSheet("""
                    QWidget {
                        background-color: #f8fafc;
                        border-radius: 12px;
                        padding: 8px;
                    }
                    QWidget:hover {
                        background-color: #f1f5f9;
                    }
                """)
                row_layout = QHBoxLayout(row_widget)
                row_layout.setContentsMargins(15, 8, 15, 8)
                
                # Create radio button group with emoji states
                button_group = QButtonGroup(row_widget)
                button_widget = QWidget()
                button_layout = QHBoxLayout(button_widget)
                button_layout.setContentsMargins(0, 0, 0, 0)
                button_layout.setSpacing(30)  # Increased spacing between buttons
                
                for i, (_, outlined, filled, color) in enumerate(ratings):
                    radio = QRadioButton()
                    radio.setStyleSheet(f"""
                        QRadioButton {{
                            font-size: 20px;
                            color: {color};
                            padding: 8px;
                            border-radius: 8px;
                        }}
                        QRadioButton::indicator {{
                            width: 0px;
                            height: 0px;
                        }}
                        QRadioButton::checked {{
                            background-color: {color}10;
                        }}
                    """)
                    # Set the text to outlined emoji by default
                    radio.setText(outlined)
                    # When checked, change to filled emoji
                    radio.toggled.connect(lambda checked, r=radio, o=outlined, f=filled: 
                                        r.setText(f if checked else o))
                    
                    button_group.addButton(radio, i)
                    self.preference_buttons[meal.name] = button_group
                    button_layout.addWidget(radio)
                
                # Add meal name with modern styling
                name_label = QLabel(meal.name)
                name_label.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
                name_label.setStyleSheet("""
                    font-size: 14px;
                    color: #1e293b;
                    font-weight: 500;
                    padding: 0 15px;
                """)
                
                row_layout.addWidget(button_widget)
                row_layout.addWidget(name_label)
                
                section_layout.addWidget(row_widget)
            
            scroll_layout.addWidget(section_group)
        
        scroll.setWidget(scroll_content)
        scroll.setStyleSheet("""
            QScrollArea {
                border: none;
                background-color: white;
            }
            QScrollBar:vertical {
                border: none;
                background: #f1f5f9;
                width: 8px;
                border-radius: 4px;
                margin: 0;
            }
            QScrollBar::handle:vertical {
                background: #94a3b8;
                border-radius: 4px;
                min-height: 30px;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0px;
            }
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {
                background: none;
            }
        """)
        preferences_layout.addWidget(scroll)
        
        # Add to main tabs
        tabs.addTab(preferences_tab, "Preferences")
        # add patient viewer tab
        patients_tab = QWidget()
        patients_layout = QVBoxLayout(patients_tab)
        self.patients_table = QTableWidget()
        self.patients_table.setColumnCount(4)
        self.patients_table.setHorizontalHeaderLabels(
        ["ID", "First Name", "Last Name", "Date of Birth"]
        )
        self.patients_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Fixed)
        self.patients_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.patients_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        self.patients_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.Fixed)
        self.patients_table.setColumnWidth(0, 50)
        self.patients_table.setColumnWidth(3, 120)
        patients_layout.addWidget(self.patients_table)
        tabs.addTab(patients_tab, "Patients")
        self.patients_table.cellClicked.connect(self.on_patient_clicked)
        self.load_patients()
        # Add tab widget to main layout
        layout.addWidget(tabs)
        
        # Initialize items list
        
        # Days of the week in Arabic
        self.days = [
            "السبت", "الأحد", "الإثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة",
            "السبت", "الأحد", "الإثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة"
        ]
        
        # Initialize the table with dropdowns
        self.initialize_table()

        # Set table properties for better Arabic text display
        self.table.setLayoutDirection(Qt.RightToLeft)
        header = self.table.horizontalHeader()
        header.setDefaultAlignment(Qt.AlignRight | Qt.AlignVCenter)
        
        # Set minimum row height for better readability
        self.table.verticalHeader().setDefaultSectionSize(50)
        self.table.setAlternatingRowColors(True)
        self.table.setShowGrid(True)
        self.table.setGridStyle(Qt.SolidLine)
        self.table.setStyleSheet(self.table.styleSheet() + """
            QTableWidget {
                gridline-color: #f3f4f6;
            }
            QTableWidget::item {
                padding: 12px;
                border-bottom: 1px solid #f3f4f6;
            }
            QTableWidget::item:alternate {
                background-color: #fafafa;
            }
        """)

        # Create admin tab
        admin_tab = QWidget()
        admin_layout = QVBoxLayout(admin_tab)
        
        # Create tab widget for admin sections
        admin_tabs = QTabWidget()
        
        # Meal Items Management Tab
        meal_items_tab = QWidget()
        meal_items_layout = QVBoxLayout(meal_items_tab)
        
        # Meal Items Table
        meal_items_group = QGroupBox("Meal Items")
        meal_items_table_layout = QVBoxLayout()
        
        self.meal_items_table = QTableWidget()
        self.meal_items_table.setColumnCount(4)
        self.meal_items_table.setHorizontalHeaderLabels(["الاسم", "وقت الوجبة", "المجموعة", "الإجراءات"])
        self.meal_items_table.horizontalHeader().setStretchLastSection(False)
        self.meal_items_table.setStyleSheet("""
            QTableWidget {
                background-color: white;
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                gridline-color: #e5e7eb;
            }
            QTableWidget::item {
                padding: 8px;
            }
            QHeaderView::section {
                background-color: #f3f4f6;
                padding: 8px;
                border: none;
                border-bottom: 2px solid #e5e7eb;
                font-weight: bold;
                color: #374151;
            }
            QPushButton {
                padding: 6px 12px;
                border-radius: 6px;
                font-size: 12px;
                min-width: 60px;
            }
        """)
        
        # Add/Edit Meal Item Form
        meal_item_form = QGroupBox("إضافة/تعديل وجبة")
        meal_item_form.setStyleSheet("""
            QGroupBox {
                font-size: 14px;
                font-weight: bold;
                margin-top: 2ex;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                subcontrol-position: top center;
                padding: 0 3px;
            }
            QLineEdit, QComboBox, QSpinBox {
                padding: 8px;
                border: 2px solid #e5e7eb;
                border-radius: 8px;
                background-color: #ffffff;
                color: #111827;
                font-size: 13px;
                min-width: 250px;
                min-height: 20px;
            }
            QLineEdit:focus, QComboBox:focus, QSpinBox:focus {
                border-color: #3b82f6;
            }
            QLabel {
                font-size: 13px;
                font-weight: 500;
                color: #374151;
                qproperty-alignment: 'AlignRight';
            }
            QPushButton {
                padding: 8px 16px;
                border-radius: 8px;
                font-weight: bold;
                font-size: 13px;
                min-width: 100px;
                min-height: 35px;
            }
            QPushButton[type="edit"] {
                background-color: #3b82f6;
                color: white;
            }
            QPushButton:hover {
                opacity: 0.9;
            }
            QPushButton:disabled {
                background-color: #9ca3af;
            }
        """)
        
        self.meal_name_edit = QLineEdit()
        self.meal_name_edit.setPlaceholderText("أدخل اسم الوجبة")
        self.meal_name_edit.setLayoutDirection(Qt.RightToLeft)
        
        self.meal_eat_time_combo = QComboBox()
        self.meal_eat_time_combo.addItems(["Breakfast", "Lunch", "Dinner"])
        self.meal_eat_time_combo.setLayoutDirection(Qt.LeftToRight)
        
        self.meal_group_spin = QSpinBox()
        self.meal_group_spin.setRange(1, 2)
        self.meal_group_spin.setLayoutDirection(Qt.LeftToRight)
        
        form_layout = QFormLayout()
        form_layout.setSpacing(15)
        form_layout.setContentsMargins(20, 30, 20, 20)
        
        form_layout.addRow("الاسم:", self.meal_name_edit)
        form_layout.addRow("وقت الوجبة:", self.meal_eat_time_combo)
        form_layout.addRow("المجموعة:", self.meal_group_spin)
        
        # Add/Edit buttons
        button_layout = QHBoxLayout()
        button_layout.setSpacing(10)
        
        self.add_meal_button = QPushButton("إضافة")
        self.add_meal_button.setProperty("type", "edit")
        self.add_meal_button.clicked.connect(self.add_meal_item)
        
        self.edit_meal_button = QPushButton("تعديل")
        self.edit_meal_button.setProperty("type", "edit")
        self.edit_meal_button.clicked.connect(self.edit_meal_item)
        self.edit_meal_button.setEnabled(False)
        
        button_layout.addWidget(self.add_meal_button)
        button_layout.addWidget(self.edit_meal_button)
        button_layout.setAlignment(Qt.AlignCenter)
        
        form_layout.addRow(button_layout)
        meal_item_form.setLayout(form_layout)

        meal_items_table_layout.addWidget(self.meal_items_table)
        meal_items_group.setLayout(meal_items_table_layout)
        
        meal_items_layout.addWidget(meal_items_group)
        meal_items_layout.addWidget(meal_item_form)
        
        # Excluded Foods Management Tab
        excluded_foods_tab = QWidget()
        excluded_foods_layout = QVBoxLayout(excluded_foods_tab)
        
        # Diabetes Excluded Foods
        diabetes_group = QGroupBox("Diabetes Excluded Foods")
        diabetes_layout = QVBoxLayout()
        
        self.diabetes_table = QTableWidget()
        self.diabetes_table.setColumnCount(2)
        self.diabetes_table.setHorizontalHeaderLabels(["الاسم", "الإجراءات"])
        self.diabetes_table.horizontalHeader().setStretchLastSection(False)
        self.diabetes_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.diabetes_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Fixed)
        self.diabetes_table.setColumnWidth(1, 150)
        self.diabetes_table.verticalHeader().setDefaultSectionSize(50)
        
        
        diabetes_layout.addWidget(self.diabetes_table)
        diabetes_group.setLayout(diabetes_layout)
        
        # Kidney Excluded Foods
        kidney_group = QGroupBox("Kidney Excluded Foods")
        kidney_layout = QVBoxLayout()
        
        self.kidney_table = QTableWidget()
        self.kidney_table.setColumnCount(2)
        self.kidney_table.setHorizontalHeaderLabels(["الاسم", "الإجراءات"])
        self.kidney_table.horizontalHeader().setStretchLastSection(False)
        self.kidney_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.kidney_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Fixed)
        self.kidney_table.setColumnWidth(1, 150)
        self.kidney_table.verticalHeader().setDefaultSectionSize(50)
        kidney_layout.addWidget(self.kidney_table)
        kidney_group.setLayout(kidney_layout)
        
        excluded_foods_layout.addWidget(diabetes_group)
        excluded_foods_layout.addWidget(kidney_group)
        
        # Add tabs to admin tabs
        admin_tabs.addTab(meal_items_tab, "Meal Items")
        admin_tabs.addTab(excluded_foods_tab, "Excluded Foods")
        
        admin_layout.addWidget(admin_tabs)
        
        # Add admin tab to main tabs
        tabs.addTab(admin_tab, "Admin")
        role = self.user.role.name.lower()
        allowed_tabs = {
            'admin':    ["Meal Planner", "Health Conditions", "Exclude Items", "Preferences", "Patients", "Admin"],
            'secretary':["Preferences", "Patients"],
            'patient':  ["Preferences"],
        }
        for i in reversed(range(tabs.count())):
            if tabs.tabText(i) not in allowed_tabs[role]:
                tabs.removeTab(i)
        if role == 'patient':
            pid = self.user.patient.id
            print(">> pid:", pid)
            self.current_patient_id = pid
            print(">> self.user:", self.user)
            print(">>   type:", type(self.user))
            print(">>   attrs:", dir(self.user))
            self.sample_input.setDisabled(True)
            self.sample_input.setText(str(pid))
            self.name_input.setText(f"{self.user.patient.first_name} {self.user.patient.last_name}")
            self.name_input.setDisabled(True)
            self.load_preferences_for_patient(pid)
        # if role == 'patient':
        #     for tab_widget_label in ("Admin","Preferences"):
        #         idx = tabs.indexOf(tabs.findChild(QWidget,tab_widget_label))
        #         if idx != -1:
        #             tabs.removeTab(idx)
        # elif role == 'secretary':
        #     idx = tabs.indexOf(admin_tab)
        #     if idx != -1:
        #         tabs.removeTab(idx)
        # Initialize tables
        self.initialize_meal_items_table()
        self.initialize_excluded_foods_tables()

    def update_health_conditions(self):
        """Update the health conditions array based on checkbox states"""
        self.health_conditions = []
        if self.healthy_checkbox.isChecked():
            self.health_conditions.append(1)
        if self.diabetes_checkbox.isChecked():
            self.health_conditions.append(2)
        if self.kidney_checkbox.isChecked():
            self.health_conditions.append(3)
        
        # Reinitialize table with updated conditions
        self.initialize_table()

    def get_excluded_items(self):
        """Get combined list of manually excluded items and condition-based exclusions"""
        #1) manually toggled ones:
        manual = [name for name, cb in self.exclusion_checkboxes.items() if cb.isChecked()]
        if not self.health_conditions:
            return manual
        cond_ids = self.health_conditions
        q = (
            self.db.query(MealItem.name)
            .join(ExclusionRule,ExclusionRule.item_id == MealItem.item_id)
            .filter(ExclusionRule.condition_id.in_(cond_ids))
        )
        db_excluded = [row[0] for row in q.all()]
        return list(set(manual + db_excluded))

    def initialize_table(self):
        self.table.setRowCount(len(self.days))
        
        # Get excluded items
        excluded_items = self.get_excluded_items()
        
        # Update ComboBox style to ensure text visibility
        combo_style = """
        QComboBox {
            padding: 10px 12px;
            border: 2px solid #e5e7eb;
            border-radius: 8px;
            background-color: #ffffff;
            color: #111827;
            min-height: 40px;
            min-width: 250px;
            font-size: 14px;
        }
        QComboBox QAbstractItemView {
            border: 1px solid #e5e7eb;
            border-radius: 8px;
            background-color: #ffffff;
            selection-background-color: #dbeafe;
            selection-color: #1e40af;
            padding: 4px;
            font-size: 13px;
        }
        QComboBox QAbstractItemView::item {
            min-height: 36px;
            padding: 6px 8px;
        }
        QComboBox::drop-down{
            width: 24px;
            border: none;
        }
        QComboBox::down-arrow{
        image: none;
        width: 0;
        height: 0;
        margin-right: 8px;
        border-left: 5px solid transparent;
        border-right: 5px solid transparent;
        border-top: 5px solid #6b7280;
        }
        """
        # # Create dropdowns for each meal cell
        for row in range(len(self.days)):
            # Day column
            day_item = QTableWidgetItem(self.days[row])
            day_item.setTextAlignment(Qt.AlignRight | Qt.AlignVCenter)
            self.table.setItem(row, 0, day_item)
            
            # Breakfast column
            breakfast_combo = QComboBox()
            breakfast_combo.setStyleSheet(combo_style)

            breakfast_combo.setLayoutDirection(Qt.RightToLeft)
            breakfast_combo.view().setLayoutDirection(Qt.RightToLeft)
            breakfast_combo.setSizeAdjustPolicy(QComboBox.AdjustToContents)
            breakfast_combo.setMinimumContentsLength(1)
            breakfast_combo.setMinimumHeight(42)
            breakfast_combo.setEditable(False)
            # breakfast_group1_items = [item["name"] for item in self.items 
            #                         if item["eat_time"] == "Breakfast" 
            #                         and item["group"] == 1
            #                         and item["name"] not in excluded_items]
            breakfast_group1_items = [
                m.name for m in (
                    self.db
                    .query(MealItem)
                    .filter_by(eat_time = 'Breakfast',group = 1)
                    .all()
                )
            ]
            breakfast_group2_items = [
                m.name for m in (
                    self.db
                    .query(MealItem)
                    .filter_by(eat_time = 'Breakfast',group = 2)
                    .all()
                )
            ]
            breakfast_combinations = [f"{g1} + {g2}"
                                  for g1 in breakfast_group1_items
                                  for g2 in breakfast_group2_items]
            breakfast_combo.addItems(breakfast_combinations)
            breakfast_combo.addItems(breakfast_combinations)
            if breakfast_combinations:
                breakfast_combo.setCurrentIndex(random.randint(0, len(breakfast_combinations) - 1))
            self.table.setCellWidget(row, 1, breakfast_combo)
            
            # Lunch column
            lunch_combo = QComboBox()
            lunch_combo.setStyleSheet(combo_style)
            lunch_combo.setLayoutDirection(Qt.RightToLeft)
            lunch_combo.view().setLayoutDirection(Qt.RightToLeft)
            lunch_combo.setSizeAdjustPolicy(QComboBox.AdjustToContents)
            lunch_combo.setMinimumContentsLength(1)
            lunch_combo.setMinimumHeight(42)
            lunch_combo.setEditable(False)
            # lunch_group1_items = [item["name"] for item in self.items 
            #                     if item["eat_time"] == "Lunch" 
            #                     and item["group"] == 1
            #                     and item["name"] not in excluded_items]
            # lunch_group2_items = [item["name"] for item in self.items 
            #                     if item["eat_time"] == "Lunch" 
            #                     and item["group"] == 2
            #                     and item["name"] not in excluded_items]
            # lunch_combinations = [f"{g1} + {g2}" for g1 in lunch_group1_items 
            #                     for g2 in lunch_group2_items]
            lunch_group1_items = [
                m.name for m in (
                    self.db
                    .query(MealItem)
                    .filter_by(eat_time='Lunch', group=1)
                    .filter(~MealItem.name.in_(excluded_items))
                    .all()
            )
            ]
            lunch_group2_items = [
            m.name for m in (
                self.db
                    .query(MealItem)
                    .filter_by(eat_time='Lunch', group=2)
                    .filter(~MealItem.name.in_(excluded_items))
                    .all()
            )
            ]
            lunch_combinations = [f"{g1} + {g2}"
                              for g1 in lunch_group1_items
                              for g2 in lunch_group2_items]
            lunch_combo.addItems(lunch_combinations)
            if lunch_combinations:
                lunch_combo.setCurrentIndex(random.randint(0, len(lunch_combinations) - 1))
            self.table.setCellWidget(row, 2, lunch_combo)
            
            # Dinner column
            dinner_combo = QComboBox()
            dinner_combo.setStyleSheet(combo_style)
            dinner_combo.setLayoutDirection(Qt.RightToLeft)
            dinner_combo.view().setLayoutDirection(Qt.RightToLeft)
            dinner_combo.setSizeAdjustPolicy(QComboBox.AdjustToContents)
            dinner_combo.setMinimumContentsLength(1)
            dinner_combo.setMinimumHeight(15)
            dinner_combo.setEditable(False)
            
            # dinner_items = [item["name"] for item in self.items 
            #               if item["eat_time"] == "Dinner" 
            #               and item["group"] == 1
            #               and item["name"] not in excluded_items]
            dinner_items = [
                m.name for m in (
                    self.db
                    .query(MealItem)
                    .filter_by(eat_time='Dinner', group=1)
                    .filter(~MealItem.name.in_(excluded_items))
                    .all()
                )
            ]
            dinner_combo.addItems(dinner_items)
            if dinner_items:
                dinner_combo.setCurrentIndex(random.randint(0, len(dinner_items) - 1))
            self.table.setCellWidget(row, 3, dinner_combo)

        # Adjust table column widths
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Fixed)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(3, QHeaderView.Stretch)
        self.table.setColumnWidth(0, 100)  # Width for the day column

        # Enable save buttons after table is initialized
        self.save_excel_button.setEnabled(True)
        self.save_word_button.setEnabled(True)
        self.save_pdf_button.setEnabled(True)
        self.save_gdocs_button.setEnabled(True)
        self.export_to_word_template_button.setEnabled(True)

    def generate_meal_plan(self):
        try:
            # 1) Read & validate your 7‑day category totals
            green_n  = self.category_a_spin.value()
            yellow_n = self.category_b_spin.value()
            red_n    = self.category_c_spin.value()
            if green_n + yellow_n + red_n != 7:
                QMessageBox.warning(self, "Error", "Category counts must sum to 7")
                return

            # 2) Figure out which patient’s prefs to use
            role = self.user.role.name.lower()
            if role == 'patient':
                pid = self.user.patient.id
            else:
                if not hasattr(self, 'current_patient_id'):
                    QMessageBox.warning(self, "Error", "Please select a patient first.")
                    return
                pid = self.current_patient_id

            # 3) Load their stored preference‐weights
            prefs = {
                p.meal_name: self.RATING_WEIGHTS.get(p.rating, 1.0)
                for p in self.db.query(Preference).filter_by(patient_id=pid)
            }

            # 4) Helpers ----------------------------------------------------------------

            def pick_by_color(pool, g, y, r):
                """
                pool: list of MealItem
                returns exactly g+y+r picks, sampling WITH replacement from each color,
                falling back to the entire pool if a color‐bin is empty.
                """
                by_col = {
                    "Green":  [m for m in pool if m.color == "Green"],
                    "Yellow": [m for m in pool if m.color == "Yellow"],
                    "Red":    [m for m in pool if m.color == "Red"],
                }
                out = []
                for col, cnt in (("Green", g), ("Yellow", y), ("Red", r)):
                    choices = by_col[col] or pool
                    # now we know choices is nonempty
                    out += random.choices(choices, k=cnt)
                return out

            def weighted_without_replacement(items, weights):
                """
                items, weights: same length lists
                returns a reordering of items, sampling WITHOUT replacement
                proportional to weights.
                """
                items = list(items)
                wts   = list(weights)
                out   = []
                while items:
                    total = sum(wts)
                    r = random.random() * total
                    cum = 0.0
                    for i, w in enumerate(wts):
                        cum += w
                        if r <= cum:
                            out.append(items.pop(i))
                            wts.pop(i)
                            break
                return out

            # 5) Build your raw meal‑item pools, excluding anything the user checked off
            excluded = set(self.get_excluded_items())
            eat = lambda tm, grp: [
                m for m in self.items
                if m.eat_time == tm and m.group == grp and m.name not in excluded
            ]
            bg1, bg2 = eat("Breakfast", 1), eat("Breakfast", 2)
            lg1, lg2 = eat("Lunch",     1), eat("Lunch",     2)
            dn       = eat("Dinner",    1)

            # 6) Decide how many “weeks” your table holds
            rows       = self.table.rowCount()
            weeks_full = rows // 7
            extra      = rows % 7

            # 7) Generate one fresh 7‑day block PER week (plus one if there’s a partial)
            b1_all = []
            b2_all = []
            l1_all = []
            l2_all = []
            d_all  = []

            for _ in range(weeks_full + (1 if extra else 0)):
                # a) pick g/y/r EXACTLY from each sub‑pool
                b1p = pick_by_color(bg1, green_n, yellow_n, red_n)
                b2p = pick_by_color(bg2, green_n, yellow_n, red_n)
                l1p = pick_by_color(lg1, green_n, yellow_n, red_n)
                l2p = pick_by_color(lg2, green_n, yellow_n, red_n)
                dp  = pick_by_color(dn,   green_n, yellow_n, red_n)

                # b) reorder by preference‐weight (no replacement)
                b1w = weighted_without_replacement(b1p, [prefs.get(m.name,1.0) for m in b1p])
                b2w = weighted_without_replacement(b2p, [prefs.get(m.name,1.0) for m in b2p])
                l1w = weighted_without_replacement(l1p, [prefs.get(m.name,1.0) for m in l1p])
                l2w = weighted_without_replacement(l2p, [prefs.get(m.name,1.0) for m in l2p])
                dw  = weighted_without_replacement(dp,  [prefs.get(m.name,1.0) for m in dp ])

                # c) stash them
                b1_all += b1w
                b2_all += b2w
                l1_all += l1w
                l2_all += l2w
                d_all  += dw

            # 8) Chop off any extra days so we have exactly “rows” entries
            b1_all = b1_all[:rows]
            b2_all = b2_all[:rows]
            l1_all = l1_all[:rows]
            l2_all = l2_all[:rows]
            d_all  = d_all[:rows]

            # 9) Write them back into your  table
            for i in range(rows):
                b1,b2 = b1_all[i], b2_all[i]
                self.table.cellWidget(i,1).setCurrentText(f"{b1.name} + {b2.name}")

                l1,l2 = l1_all[i], l2_all[i]
                self.table.cellWidget(i,2).setCurrentText(f"{l1.name} + {l2.name}")

                d = d_all[i]
                self.table.cellWidget(i,3).setCurrentText(d.name)

            # 10) Enable your export buttons
            for btn in (
                self.save_excel_button,
                self.save_word_button,
                self.save_pdf_button,
                self.save_gdocs_button,
                self.export_to_word_template_button
            ):
                btn.setEnabled(True)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate meal plan: {e}")

    def save_to_excel(self):
        try:
            file_name,_ = QFileDialog.getSaveFileName(self, "Save Excel File", "", "Excel Files (*.xlsx)")
            if not file_name:
                return
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.title = "خطة الوجبات"
            # 1) Patient info
            sheet.cell(row=1, column=1).value = "Name:"
            sheet.cell(row=1, column=2).value = self.name_input.text()
            print(">> self.name_input.text():", self.name_input.text())
            sheet.cell(row=2, column=1).value = "Sample Number:"
            sheet.cell(row=2, column=2).value = self.sample_input.text()
            ## 2) Health Conditions
            sheet.cell(row=1, column=4).value = "Health Conditions:"
            conditions = []
            if 1 in self.health_conditions:
                conditions.append("Healthy")
            if 2 in self.health_conditions:
                conditions.append("Diabetes")
            if 3 in self.health_conditions:
                conditions.append("Kidney Disease")
            sheet.cell(row=1, column=5).value = ", ".join(conditions)
            headers = ["اليوم", "الإفطار", "الغداء", "العشاء"]
            for col, header in enumerate(headers, 1):
                sheet.cell(row=3, column=col).value = header
            excluded = set(self.get_excluded_items())
            def pool(time, grp):
                return [
                m.name for m in self.items
                if m.eat_time == time
                   and m.group == grp
                   and m.name not in excluded
                ]

            bf1 = pool("Breakfast", 1)
            bf2 = pool("Breakfast", 2)
            ln1 = pool("Lunch",     1)
            ln2 = pool("Lunch",     2)
            dn  = pool("Dinner",    1)

        # 4) Write the hidden dropdown sources into F/H columns
        #    (same as before, but variable names updated)
            breakfast_combinations = [f"{a} + {b}" for a in bf1 for b in bf2]
            lunch_combinations     = [f"{a} + {b}" for a in ln1 for b in ln2]

        # write them into columns F, G, H and capture ranges
            for i, combo in enumerate(breakfast_combinations, start=1):
                sheet.cell(row=i, column=6).value = combo
            breakfast_range = f"$F$1:$F${len(breakfast_combinations)}"

            for i, combo in enumerate(lunch_combinations, start=1):
                sheet.cell(row=i, column=7).value = combo
            lunch_range = f"$G$1:$G${len(lunch_combinations)}"

            for i, item in enumerate(dn, start=1):
                sheet.cell(row=i, column=8).value = item
            dinner_range = f"$H$1:$H${len(dn)}"

        # 5) Now write your 7-day plan (rows 4→10) and attach DataValidation
            for r in range(self.table.rowCount()):
                excel_row = r + 4
            # Day name
                sheet.cell(row=excel_row, column=1).value = self.table.item(r, 0).text()

            # Breakfast, lunch, dinner texts
                bf = self.table.cellWidget(r, 1).currentText()
                ln = self.table.cellWidget(r, 2).currentText()
                dn = self.table.cellWidget(r, 3).currentText()

                sheet.cell(row=excel_row, column=2).value = bf
                sheet.cell(row=excel_row, column=3).value = ln
                sheet.cell(row=excel_row, column=4).value = dn

                # Attach dropdown validations
                dv_bf = DataValidation(type="list", formula1=f"={breakfast_range}", allow_blank=True)
                dv_ln = DataValidation(type="list", formula1=f"={lunch_range}",     allow_blank=True)
                dv_dn = DataValidation(type="list", formula1=f"={dinner_range}",    allow_blank=True)

                sheet.add_data_validation(dv_bf)
                sheet.add_data_validation(dv_ln)
                sheet.add_data_validation(dv_dn)

                dv_bf.add(sheet.cell(row=excel_row, column=2))
                dv_ln.add(sheet.cell(row=excel_row, column=3))
                dv_dn.add(sheet.cell(row=excel_row, column=4))

        # 6) Hide the helper columns
            sheet.column_dimensions['F'].hidden = True
            sheet.column_dimensions['G'].hidden = True
            sheet.column_dimensions['H'].hidden = True

            workbook.save(file_name)
            QMessageBox.information(self, "Success", "Meal plan saved successfully!")
        except Exception as e:
            QMessageBox.critical(self, "Error", str(e))
    def save_to_word(self):
        try:
            file_name, _ = QFileDialog.getSaveFileName(
                self, "Save Word File", "", "Word Files (*.docx)"
            )
            if not file_name:
                return

            # 1) Figure out which names to exclude
            excluded = set(self.get_excluded_items())

            # 2) Build pools of names using ORM attributes
            def pool(time, grp):
                return [
                    m.name for m in self.items
                    if m.eat_time == time
                       and m.group == grp
                       and m.name not in excluded
                ]

            bf1 = pool("Breakfast", 1)
            bf2 = pool("Breakfast", 2)
            ln1 = pool("Lunch",     1)
            ln2 = pool("Lunch",     2)
            dn  = pool("Dinner",    1)

            breakfast_combinations = [f"{a} + {b}" for a in bf1 for b in bf2]
            lunch_combinations     = [f"{a} + {b}" for a in ln1 for b in ln2]
            dinner_items           = dn

            # 3) Create the document
            doc = Document()
            conditions = []
            if 1 in self.health_conditions:
                conditions.append("Healthy")
            if 2 in self.health_conditions:
                conditions.append("Diabetes")
            if 3 in self.health_conditions:
                conditions.append("Kidney Disease")
            doc.add_paragraph(f"Health Conditions: {', '.join(conditions)}")

            # 4) Build a 1+14 row table
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(["اليوم", "الإفطار", "الغداء", "العشاء"]):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # helper to insert a Word dropdown in a cell
            def create_dropdown(cell, options, selected):
                try:
                    sdt = OxmlElement('w:sdt')
                    sdtPr = OxmlElement('w:sdtPr')
                    ddl = OxmlElement('w:dropDownList')
                    for opt in options:
                        li = OxmlElement('w:listItem')
                        li.set(qn('w:displayText'), opt)
                        li.set(qn('w:value'), opt)
                        ddl.append(li)
                    sdtPr.append(ddl)
                    sdt.append(sdtPr)

                    content = OxmlElement('w:sdtContent')
                    p = OxmlElement('w:p')
                    r = OxmlElement('w:r')
                    t = OxmlElement('w:t')
                    t.text = selected or (options[0] if options else "")
                    r.append(t)
                    p.append(r)
                    content.append(p)
                    sdt.append(content)

                    cell._element.clear_content()
                    cell._element.append(sdt)
                except Exception:
                    cell.text = selected or (options[0] if options else "")

            # 5) Fill in each of the 14 rows
            for row_idx in range(self.table.rowCount()):
                cells = table.add_row().cells
                # Day name
                day = self.table.item(row_idx, 0).text()
                cells[0].text = day
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # Grab what the user has currently selected in the UI table
                bf_sel = self.table.cellWidget(row_idx, 1).currentText()
                ln_sel = self.table.cellWidget(row_idx, 2).currentText()
                dn_sel = self.table.cellWidget(row_idx, 3).currentText()

                # Insert dropdowns
                create_dropdown(cells[1], breakfast_combinations, bf_sel)
                create_dropdown(cells[2], lunch_combinations,     ln_sel)
                create_dropdown(cells[3], dinner_items,           dn_sel)

            # 6) Save
            doc.save(file_name)
            QMessageBox.information(self, "Success", "Meal plan saved successfully!")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save document: {e}")
    def save_to_pdf(self):
        try:
        # Get save file path
            file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save PDF Document",
            "",
            "PDF Files (*.pdf)"
        )
        
            if not file_path:
                return

            from weasyprint import HTML, CSS
            from weasyprint.text.fonts import FontConfiguration
        
        # Create HTML content
            html_content = f"""
        <html dir="rtl">
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 1cm;
                }}
                body {{
                    font-family: Arial, sans-serif;
                    direction: rtl;
                }}
                h1 {{
                    text-align: center;
                    color: black;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin-top: 20px;
                }}
                th, td {{
                    border: 1px solid black;
                    padding: 8px;
                    text-align: center;
                }}
                th {{
                    background-color: #f2f2f2;
                    font-weight: bold;
                }}
            </style>
        </head>
        <body>
            <h1>خطة الوجبات الأسبوعية</h1>
            <table>
                <tr>
                    <th>اليوم</th>
                    <th>الإفطار</th>
                    <th>الغداء</th>
                    <th>العشاء</th>
                </tr>
        """
        
        # Add table rows
            for row in range(self.table.rowCount()):
                day = self.table.item(row, 0).text()
                breakfast = self.table.cellWidget(row, 1).currentText()
                lunch = self.table.cellWidget(row, 2).currentText()
                dinner = self.table.cellWidget(row, 3).currentText()
            
                html_content += f"""
                <tr>
                    <td>{day}</td>
                    <td>{breakfast}</td>
                    <td>{lunch}</td>
                    <td>{dinner}</td>
                </tr>
            """
        
            html_content += """
            </table>
        </body>
        </html>
        """
        
        # Configure fonts
            font_config = FontConfiguration()
        
        # Create PDF
            HTML(string=html_content).write_pdf(
            file_path,
            font_config=font_config,
            presentational_hints=True
            )
        
            QMessageBox.information(self, "Success", "Meal plan saved to PDF successfully!")
        
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save PDF: {str(e)}")
    def get_breakfast_combinations(self):
        excluded  = set(self.get_excluded_items())
        group1 = [m.name for m in self.items
                  if m.eat_time  == "Breakfast" and m.group == 1 and m.name not in excluded]
        group2 = [m.name for m in self.items
                  if m.eat_time  == "Breakfast" and m.group == 2 and m.name not in excluded]
        return [f"{a} + {b}" for a in group1 for b in group2]
    def get_lunch_combinations(self):
        excluded = set(self.get_excluded_items())
        group1 = [m.name for m in self.items
                  if m.eat_time  == "Lunch" and m.group == 1 and m.name not in excluded]
        group2 = [m.name for m in self.items
                  if m.eat_time  == "Lunch" and m.group == 2 and m.name not in excluded]
        return [f"{a} + {b}" for a in group1 for b in group2]
    def get_dinner_items(self):
        excluded = set(self.get_excluded_items())
        return [m.name for m in self.items
                if m.eat_time == "Dinner" and m.group == 1 and m.name not in excluded]
        
    def get_gdocs_service(self):
        SCOPES = ['https://www.googleapis.com/auth/documents', 'https://www.googleapis.com/auth/drive.file']
        flow = InstalledAppFlow.from_client_secrets_file("google_doc_credential.json", SCOPES)
        creds = flow.run_local_server(port=0)
        return build('docs', 'v1', credentials=creds)
    def get_health_conditions_text(self):
        conditions = []
        if hasattr(self,"health_conditions"):
            if 1 in self.health_conditions:
                conditions.append("Healthy")
            if 2 in self.health_conditions:
                conditions.append("Diabetes")
            if 3 in self.health_conditions:
                conditions.append("Kidney Disease")
        return conditions;
    def save_to_gdoc(self):
        try:
            # SECTION: Authentication and Document Creation
            print("SECTION: Authentication and Document Creation")
            SCOPES = ['https://www.googleapis.com/auth/documents', 'https://www.googleapis.com/auth/drive.file']
            if not os.path.exists("google_doc_credential.json"):
                raise FileNotFoundError("Google Docs credentials file 'google_doc_credential.json' not found.")

            from google_auth_oauthlib.flow import InstalledAppFlow
            from googleapiclient.discovery import build
            flow = InstalledAppFlow.from_client_secrets_file("google_doc_credential.json", SCOPES)
            creds = flow.run_local_server(port=0)
            service = build('docs', 'v1', credentials=creds)

            title = "خطة الوجبات الأسبوعية"
            doc = service.documents().create(body={"title": title}).execute()
            doc_id = doc['documentId']
            print(f"Created document with ID: {doc_id}")
            print("SECTION: INSERT TABLE ")
            table_rows  = self.table.rowCount() + 1
            table_columns = 4
            table_start_index = 1
            requests = []
            requests.append({
                'insertTable':{
                    'rows':table_rows,
                    'columns':table_columns,
                    'location':{'index':table_start_index}
                }
            })
            service.documents().batchUpdate(documentId = doc_id,body={'requests':requests}).execute()
            print("Table printed successfully")
            print("SECTION FETCH CELL INDICES")
            doc_content = service.documents().get(documentId=doc_id).execute()
            cell_indices = {}
            for element in doc_content.get('body',{}).get('content',[]):
                if 'table' in element:
                    table = element['table']
                    table_start = element['startIndex']
                    print(f"Table found at index: {table_start}")
                    for r, row in enumerate(table.get('tableRows',[])):
                        for c, cell in enumerate(row.get('tableCells',[])):
                            if 'content' in cell and cell['content']:
                                para = cell['content'][0]
                                idx = para.get('startIndex')
                                if idx is not None:
                                    cell_indices[(r,c)] = idx
                                    print(f"  Cell ({r}, {c}) → startIndex={idx}")
                    break;
            print("Done fetching cell indices")
            ## SECTION: Insert Headers
            print("SECTION: INSERT HEADERS")
            headers = ["اليوم", "الإفطار", "الغداء", "العشاء"]
            cols_ordered = sorted(range(len(headers)),
                                  key = lambda c: cell_indices[(0,c)],
                                  reverse = True)
            header_requests = []
            for col in cols_ordered:
                title = headers[col]
                idx = cell_indices[(0,col)]
                print(f"-> inserting {title!r} at index {idx}")
                header_requests.append({
                    'insertText':{
                        'location':{'index':idx},
                        'text':title
                    }
                })
                # style it bold
                header_requests.append({
                    'updateTextStyle':{
                        'range':{'startIndex':idx,'endIndex':idx + len(title)},
                        'textStyle':{'bold':True},
                        'fields':'bold'
                    }
                })
                # center align
                header_requests.append({
                    'updateParagraphStyle':{
                        'range':{'startIndex':idx,'endIndex':idx + len(title)},
                        'paragraphStyle': {'alignment': 'CENTER'},
                        'fields':'alignment'
                    }
                })
            service.documents().batchUpdate(
                documentId = doc_id,
                body = {'requests':header_requests}
            ).execute()
            # NEW BLOCK: re-fetch cell_indices after headers --
            print("SECTION: REFERESH CELL INDICIES")
            doc_content = service.documents().get(documentId = doc_id).execute()
            cell_indices = {}
            for element in doc_content.get('body',{}).get('content',[]):
                if 'table' in element:
                    for r, row in enumerate(element['table']['tableRows']):
                        for c, cell in enumerate(row['tableCells']):
                            if 'content' in cell and cell['content']:
                                para = cell['content'][0]
                                idx = para.get('startIndex')
                                if idx is not None:
                                    cell_indices[(r,c)] = idx
                    break
            print("Refreshed cell indices:",cell_indices)
            # ── END NEW BLOCK ──
             # Now build your data_requests exactly as before, using this fresh `cell_indices`
            ## all good so far now the debugging will be in insert cells
            # SECTION INSERT 14 DAY ROW DATA
            print("SECTION: INSERT ROW DATA")
            days_ar = ["السبت", "الأحد", "الإثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة"]
            day_list = days_ar * 2
            data_requests = []
            for r in range(14):
                row_vals = [
                    day_list[r],
                    self.table.cellWidget(r,1).currentText(),
                    self.table.cellWidget(r, 2).currentText(),
                    self.table.cellWidget(r, 3).currentText()
                ]
                for c, text in enumerate(row_vals):
                    idx = cell_indices[(r+1,c)]
                    payload = text + "\n"
                    # insert text+newline
                    data_requests.append({
                        'insertText':{
                            'location': {'index':idx},
                            'text':payload
                        }
                    })
                    # Style the text part (not new line)
                    data_requests.append({
                        'updateTextStyle':{
                            'range':{'startIndex':idx,'endIndex':idx + len(text)},
                            'textStyle':{'fontSize':{'magnitude': 11,'unit':'PT'}},
                            'fields':'fontSize'
                        }
                    })
                    ## 3) align the paragraph
                    # Right-align the paragraph
                    data_requests.append({
                        'updateParagraphStyle':{
                            'range':{'startIndex':idx,'endIndex':idx + len(payload)},
                            'paragraphStyle':{'alignment':'END'},
                            'fields':'alignment'
                        }
                    })
            data_requests.sort(key=lambda r:(
                r.get('insertText',{}).get('location',{}).get('index',
                r.get('updateTextStyle',{}).get('range',{}).get('startIndex',0))
            ),reverse=True)
            service.documents().batchUpdate(
                documentId = doc_id,
                body={'requests':data_requests}
            ).execute()
            print("14 day data (with newline) inserted")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save to Google Docs: {str(e)}")
            import traceback
            print(traceback.format_exc())


    def save_to_template_word(self):
        try:
            # Step 1: Ask user to choose template source
            template_sources = ["Genaty Templates", "Airtable Templates"]
            source, ok = QInputDialog.getItem(
                self,
                "Select Template Source",
                "Choose template source:",
                template_sources,
                0,
                False
            )
            
            if not ok or not source:
                return
                
            # Step 2: Get templates from selected source
            if source == "Genaty Templates":
                templates_dir = "templates"
            else:  # Airtable Templates
                templates_dir = "airtable_templates"
                
            if not os.path.exists(templates_dir):
                QMessageBox.critical(self, "Error", f"Template directory '{templates_dir}' not found!")
                return
                
            template_files = [f for f in os.listdir(templates_dir) if f.endswith('.docx')]
            if not template_files:
                QMessageBox.critical(self, "Error", f"No template files found in the {source} folder!")
                return
                
            # Step 3: Show popup for user to select a template
            template_file, ok = QInputDialog.getItem(
                self,
                "Select Template",
                f"Choose a template from {source}:",
                template_files,
                0,
                False
            )
            
            if not ok or not template_file:
                return
                
            template_path = os.path.join(templates_dir, template_file)
            if not os.path.exists(template_path):
                QMessageBox.critical(self, "Error", "Template file not found!")
                return
                
            # Step 4: Get meal options and current selections
            breakfast_options = self.get_breakfast_combinations()
            lunch_options = self.get_lunch_combinations()
            dinner_options = self.get_dinner_items()
            
            # Step 5: Create replacements dictionary with current selections
            replacements = {}
            for row in range(self.table.rowCount()):
                day_num = row + 1
                breakfast_key = f"Bf{day_num}"
                lunch_key = f"Lunch{day_num}"
                dinner_key = f"Dinner{day_num}"
                breakfast_selected = self.table.cellWidget(row, 1).currentText()
                lunch_selected = self.table.cellWidget(row, 2).currentText()
                dinner_selected = self.table.cellWidget(row, 3).currentText()
                replacements[breakfast_key] = (breakfast_options, breakfast_selected)
                replacements[lunch_key] = (lunch_options, lunch_selected)
                replacements[dinner_key] = (dinner_options, dinner_selected)
                
            # Step 6: Process the template
            doc = Document(template_path)
            
            # Process shapes in the document
            for shape in doc.part.package.parts:
                if hasattr(shape, '_element') and hasattr(shape._element, 'txbx'):
                    try:
                        text_frame = shape._element.txbx
                        if text_frame is not None:
                            text_content = text_frame.text
                            for key, (options, selected) in replacements.items():
                                if key in text_content:
                                    new_p = OxmlElement('w:p')
                                    dropdown = create_dropdown_element(options, selected)
                                    new_p.append(dropdown)
                                    text_frame.clear_content()
                                    text_frame._element.append(new_p)
                    except Exception as shape_error:
                        print(f"Error processing shape: {shape_error}")
                        continue
                        
            # Alternative approach using word._element
            body = doc._element.body
            for shape in body.findall('.//w:txbxContent', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                for paragraph in shape.findall('.//w:p', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    text_content = ""
                    for run in paragraph.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        if run.text:
                            text_content += run.text
                    for key, (options, selected) in replacements.items():
                        if key in text_content:
                            for child in list(paragraph):
                                paragraph.remove(child)
                            dropdown = create_dropdown_element(options, selected)
                            paragraph.append(dropdown)
                            
            # Step 7: Save the document
            save_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save File",
                "",
                "Word Files (*.docx)"
            )
            
            if save_path:
                doc.save(save_path)
                QMessageBox.information(self, "Success", "Meal plan saved with template successfully!")
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save with template: {str(e)}")
            import traceback
            print(traceback.format_exc())

    def initialize_meal_items_table(self):
        # pull all MealItem rows and keep them in order
        self.all_meal_items = (
            self.db
            .query(MealItem)
            .order_by(MealItem.name)
            .all()
        )
        button_style = """
            QPushButton {
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 13px;
                min-width: 70px;
                min-height: 30px;
                margin: 2px;
            }
            QPushButton[type="edit"] {
                background-color: #3b82f6;
                color: white;
            }
            QPushButton[type="delete"] {
                background-color: #ef4444;
                color: white;
            }
            QPushButton:hover {
                opacity: 0.9;
            }
        """
        self.meal_items_table.setRowCount(len(self.all_meal_items))
        for row, mi in enumerate(self.all_meal_items):
            name_item = QTableWidgetItem(mi.name)
            name_item.setTextAlignment(Qt.AlignRight|Qt.AlignVCenter)
            self.meal_items_table.setItem(row,0,name_item)
            et = QTableWidgetItem(mi.eat_time)
            et.setTextAlignment(Qt.AlignCenter)
            self.meal_items_table.setItem(row,1,et)
            grp = QTableWidgetItem(str(mi.group)) 
            grp.setTextAlignment(Qt.AlignCenter)
            self.meal_items_table.setItem(row,2,grp)
            btn_edit = QPushButton("تعديل")
            btn_edit.setProperty("type","edit")
            btn_edit.setStyleSheet(button_style)
            btn_edit.clicked.connect(lambda checked, r = row: self.edit_meal_item_clicked(r))
            btn_delete = QPushButton("حذف")
            btn_delete.setProperty("type","delete")
            btn_delete.clicked.connect(lambda _, r=row: self.delete_meal_item_db(r))
            w= QWidget()
            lay=QHBoxLayout(w)
            lay.setContentsMargins(0,0,0,0)
            lay.setSpacing(4)
            lay.addWidget(btn_edit)
            lay.addWidget(btn_delete)
            self.meal_items_table.setCellWidget(row,3,w)            
        # Disable the last section resize mode to prevent extra column
        self.meal_items_table.horizontalHeader().setStretchLastSection(False)
        
        # Set column widths and resize modes
        self.meal_items_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.meal_items_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Fixed)
        self.meal_items_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Fixed)
        self.meal_items_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.Fixed)
        
        self.meal_items_table.setColumnWidth(1, 100)
        self.meal_items_table.setColumnWidth(2, 80)
        self.meal_items_table.setColumnWidth(3, 200)  # Wider for buttons
        
        # Set consistent row height
        self.meal_items_table.verticalHeader().setDefaultSectionSize(50)
        
        # Update button and table styles

        
        

    def initialize_excluded_foods_tables(self):
        # Common table style
        table_style = """
            QTableWidget {
                background-color: white;
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                gridline-color: #e5e7eb;
            }
            QTableWidget::item {
                padding: 8px;
                border-bottom: 1px solid #e5e7eb;
            }
            QHeaderView::section {
                background-color: #f3f4f6;
                padding: 8px;
                border: none;
                border-bottom: 2px solid #e5e7eb;
                font-weight: bold;
                color: #374151;
            }
        """

        # Common button style
        button_style = """
            QPushButton {
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 13px;
                min-width: 70px;
                min-height: 30px;
                margin: 2px;
            }
            QPushButton[type="delete"] {
                background-color: #ef4444;
                color: white;
            }
            QPushButton:hover {
                opacity: 0.9;
            }
        """
        # Initialize Diabetes Excluded Foods table
        self.diabetes_table.setStyleSheet(table_style)
        self.diabetes_table.setColumnCount(2)
        self.diabetes_table.setHorizontalHeaderLabels(["الاسم", "الإجراءات"])
        self.diabetes_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.diabetes_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Fixed)
        self.diabetes_table.setColumnWidth(1, 150)
        self.diabetes_table.verticalHeader().setDefaultSectionSize(50)
        # Kidney table styling & headers (mirror diabetes)
        self.kidney_table.setStyleSheet(table_style)
        self.kidney_table.setColumnCount(2)
        self.kidney_table.setHorizontalHeaderLabels(["الاسم", "الإجراءات"])
        self.kidney_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.kidney_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Fixed)
        self.kidney_table.setColumnWidth(1, 150)
        self.kidney_table.verticalHeader().setDefaultSectionSize(50)
        def populate(table,condition_id,delete_slot):
            rules = (
                self.db
                .query(ExclusionRule)
                .filter_by(condition_id = condition_id)
                .join(MealItem)
                .all()
            )
            table.setRowCount(len(rules))
            for row, rule in enumerate(rules):
                name_item = QTableWidgetItem(rule.item.name)
                name_item.setTextAlignment(Qt.AlignRight|Qt.AlignVCenter)
                table.setItem(row,0,name_item)
                btn = QPushButton("حذف")
                btn.setProperty("type","delete")
                btn.setStyleSheet(button_style)
                btn.clicked.connect(lambda _, rid=rule.rule_id: delete_slot(rid))
                w = QWidget()
                lay = QHBoxLayout(w)
                lay.setContentsMargins(0,0,0,0)
                lay.addWidget(btn)
                table.setCellWidget(row,1,w)
        populate(self.diabetes_table,2,self.delete_exclusion_rule)
        populate(self.kidney_table,3,self.delete_exclusion_rule)
        self.diabetes_table.resizeColumnsToContents()
        self.kidney_table.resizeColumnsToContents()
        
    def delete_exclusion_rule(self,rule_id):
        """Remove an Exclusion by its PK, then refresh both tables."""
        rule = self.db.get(ExclusionRule,rule_id)
        if rule:
            self.db.delete(rule)
            self.db.commit()
            self.initialize_excluded_foods_tables()
    def add_meal_item(self):
        name = self.meal_name_edit.text().strip()
        eat_time = self.meal_eat_time_combo.currentText()
        group_ = self.meal_group_spin.value()
        
        if not name:
            QMessageBox.warning(self, "Error", "Please enter a name for the meal item.")
            return
        
        # Add to MEAL_ITEMS
        mi = MealItem(name = name,eat_time = eat_time,group = group_)
        try:
            self.db.add(mi)
            self.db.commit()
        except IntegrityError:
            self.db.rollback()
            QMessageBox.warning(self,"Error","That meal already Exists")
        self.initialize_meal_items_table()
        self.meal_name_edit.clear()
        self.meal_eat_time_combo.setCurrentIndex(0)
        self.meal_group_spin.setValue(1)

    def edit_meal_item_clicked(self, row):
        """Handle edit button click in table"""
        mi = self.all_meal_items[row]
        self.current_edit_item = mi
        self.meal_name_edit.setText(mi.name)
        self.meal_eat_time_combo.setCurrentText(mi.eat_time)
        self.meal_group_spin.setValue(mi.group)
        self.edit_meal_button.setEnabled(True)
        self.add_meal_button.setEnabled(False)
        

    def edit_meal_item(self):
        mi = getattr(self,"current_edit_item",None)
        if not mi:
            return
        name = self.meal_name_edit.text().strip()
        eat_time = self.meal_eat_time_combo.currentText()
        group_ = self.meal_group_spin.value()
        if not name:
            QMessageBox.warning(self,"Error","Please enter a name")
            return
        mi.name = name
        mi.eat_time = eat_time
        mi.group = group_
        self.db.commit()
        self.initialize_meal_items_table()
        self.meal_name_edit.clear()
        self.meal_eat_time_combo.setCurrentIndex(0)
        self.meal_group_spin.setValue(1)
        self.edit_meal_button.setEnabled(False)
        self.add_meal_button.setEnabled(True)
        del self.current_edit_item
    def delete_meal_item_db(self, row):
        """"Delete the MealItem at the given row from the DB and refresh the table."""
        mi = self.all_meal_items[row]
        confirm = QMessageBox.question(
            self,
            "Confirm Delete",
            f"Are you sure you want to delete {mi.name} ?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No,
        )
        if confirm == QMessageBox.Yes:
            self.db.delete(mi)
            self.db.commit()
            self.initialize_meal_items_table()
    def add_diabetes_exclusion(self):
        name = self.diabetest_name_edit.text().strip()
        mi = self.db.query(MealItem).filter_by(name = name).first()
        if not mi:
            QMessageBox.warning(self,"Error","That meal item does not exist")
            return
        er = ExclusionRule(item_id = mi.item_id,condition_id = 2)
        self.db.add(er)
        self.db.commit()
        self.initialize_excluded_foods_tables()
        self.diabetest_name_edit.clear()

    

    def add_kidney_exclusion(self):
        name = self.kidney_name_edit.text().strip()
        mi = self.db.query(MealItem).filter_by(name = name).first()
        if not mi:
            QMessageBox.warning(self,"Error","That meal item does not exist")
            return
        er = ExclusionRule(item_id = mi.item_id,condition_id = 3)
        self.db.add(er)
        self.db.commit()
        self.initialize_excluded_foods_tables()
        self.kidney_name_edit.clear()
    def load_patients(self):
        profiles = self.db.query(PatientProfile).order_by(PatientProfile.last_name).all()
        self.patients_table.setRowCount(len(profiles))
        for row, p in enumerate(profiles):
        # store the patient_id on the QTableWidgetItem for easy retrieval
            item_id = QTableWidgetItem(str(p.id))
            item_id.setData(Qt.UserRole, p.id)
            self.patients_table.setItem(row, 0, item_id)
            self.patients_table.setItem(row, 1, QTableWidgetItem(p.first_name))
            self.patients_table.setItem(row, 2, QTableWidgetItem(p.last_name))
            self.patients_table.setItem(row, 3, QTableWidgetItem(p.dob.strftime("%Y-%m-%d")))
    
    def on_patient_clicked(self,row,column):
        item = self.patients_table.item(row,0)
        if not item:
            return
        self.current_patient_id = item.data(Qt.UserRole)
        profile = (self.db.query(PatientProfile).filter_by(id = self.current_patient_id).one_or_none())
        self.name_input.setText(f"{profile.first_name} {profile.last_name}")
        self.sample_input.setText(str(self.current_patient_id))
        self.load_preferences_for_patient(self.current_patient_id)
    ## save preferences
    def save_preferences(self):
        try:
            
            if self.user.role.name.lower() == "patient":

                patient_id = self.user.patient.id
            else:
                if not hasattr(self, "current_patient_id"):
                    QMessageBox.warning(self, "Warning", "No patient selected.")
                    return
                patient_id = self.current_patient_id
           # wipe out old prefs for that patient
            self.db.query(Preference).filter_by(patient_id=patient_id).delete()
            
            print(f"🔹 Saving Preferences for patient ID: {self.current_patient_id}")
            for meal_name, button_group in self.preference_buttons.items():
                checked = button_group.checkedButton()
                if checked:
                    idx = button_group.id(checked)
                    rating_label = self.RATING_LABELS[idx]
                else:
                    rating_label = "Not Rated"

                pref = Preference(
                    patient_id = patient_id,
                    meal_name  = meal_name,
                    rating     = rating_label
                )
                self.db.add(pref)
            self.db.commit()
            QMessageBox.information(self, "Success", "Preferences saved (printed to console).")

        except Exception as e:
            self.db.rollback()
            QMessageBox.critical(self, "Error", f"Failed to save preferences: {str(e)}")
    def load_preferences_for_patient(self,patient_id):
        prefs = self.db.query(Preference).filter_by(patient_id=patient_id).all()
        print("prefs",prefs)
        # for group in self.preference_buttons.values():
        #     for btn in group.buttons():
        #         btn.setAutoExclusive(False)
        #         btn.setChecked(False)
        #     for btn in group.buttons():
        #         btn.setAutoExclusive(True)
        for group in self.preference_buttons.values():
            group.setExclusive(False)
            for btn in group.buttons():
                btn.setChecked(False)
            group.setExclusive(True)
        for pref in prefs:
           group = self.preference_buttons.get(pref.meal_name)
           if not group:
                continue
           idx = self.rating_map.get(pref.rating)
           if idx is not None:
               btn = group.button(idx)
               if btn:
                   btn.setChecked(True)
               
                
    def setup_sample_input_listener(self):
         self.sample_input.returnPressed.connect(self.on_sample_input_entered)    
    def on_sample_input_entered(self):
        try:
            patient_id_text = self.sample_input.text().strip()
            if not patient_id_text.isdigit():
               QMessageBox.warning(self, "Invalid Input", "Please enter a valid numeric patient ID.")
               return
            patient_id = int(patient_id_text)
            profile = self.db.query(PatientProfile).filter_by(id = patient_id).first()
            if not profile:
                QMessageBox.warning(self, "Not Found", f"No patient found with ID: {patient_id}")
                return
            self.current_patient_id = patient_id
            self.name_input.setText(f"{profile.first_name} {profile.last_name}")
            print(f"here is patient id to load preferences:{patient_id}")
            self.load_preferences_for_patient(patient_id)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred: {str(e)}")
if __name__ == "__main__":
    from sqlalchemy.orm import scoped_session
    from database import SessionLocal
    app = QApplication(sys.argv)
    session = scoped_session(SessionLocal)
    dlg = LoginDialog(session)
    if not dlg.exec():
        sys.exit(0)
    me = dlg.current_user
    window = MealPlanner(session,current_user=me)
    window.show()
    sys.exit(app.exec()) 